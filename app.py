from flask import Flask, request, send_file
from flask_cors import CORS
import pypandoc
import os
import tempfile
import re
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK

# This forces Render to download the core Pandoc software
pypandoc.download_pandoc()

app = Flask(__name__)
CORS(app)

@app.route('/')
def home():
    return "Omni-AI Word Conversion Service is running!"

@app.route('/convert-to-word', methods=['POST'])
def convert_to_word():
    data = request.json
    text = data.get('text', '')
    
    # Normalize line endings to prevent splitting errors
    text = text.replace('\r\n', '\n')
    
    # --- THE OMNI-AI CLEANUP SCRIPT ---
    
    # 1. Rescue math formulas trapped inside hidden JSON widgets
    text = re.sub(r'genui[^]*"content":\s*"([^"]+)"[^]*', r'$$\n\1\n$$', text)
    
    # 2. Delete any leftover invisible tracking artifacts
    text = re.sub(r'[^]+', '', text)
    
    # 3. Fix broken Display Math from bad copy-pasting
    text = re.sub(r'^\[\s*$', '$$', text, flags=re.MULTILINE)
    text = re.sub(r'^\]\s*$', '$$', text, flags=re.MULTILINE)
    
    # 4. Fix standard Markdown Display Math
    text = text.replace(r'\[', '$$').replace(r'\]', '$$')
    
    # 5. Fix standard Markdown Inline Math
    text = text.replace(r'\(', '$').replace(r'\)', '$')

    # --- HTML TO MARKDOWN FAILSAFE ---
    # Convert standard HTML tags to Markdown so Pandoc doesn't strip them
    text = re.sub(r'<h1[^>]*>(.*?)</h1>', r'# \1\n', text, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r'<h2[^>]*>(.*?)</h2>', r'## \1\n', text, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r'<h3[^>]*>(.*?)</h3>', r'### \1\n', text, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r'<strong[^>]*>(.*?)</strong>', r'**\1**', text, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r'<b[^>]*>(.*?)</b>', r'**\1**', text, flags=re.IGNORECASE | re.DOTALL)

    # --- PAGE BREAK SYSTEM ---
    # Convert HTML page breaks into an invisible marker
    text = re.sub(r'<div\s+style="[^"]*page-break-after:\s*always;?[^"]*">\s*</div>', '\n\nOMNI_PAGE_BREAK\n\n', text, flags=re.IGNORECASE)

    # --- THE ALIGNMENT MARKER SYSTEM ---
    # Convert HTML alignment divs into invisible markers that Python will read later
    def align_replacer(match):
        alignment = match.group(1).upper()
        content = match.group(2).strip()
        return f"\n\nOMNI_ALIGN_{alignment}_START\n\n{content}\n\nOMNI_ALIGN_END\n\n"

    # Catch <div align="center"> or <div align="right">
    text = re.sub(r'<div\s+align="(center|right|left|justify)">\s*(.*?)\s*</div>', align_replacer, text, flags=re.IGNORECASE | re.DOTALL)
    
    # Catch <div style="text-align: center;">
    text = re.sub(r'<div\s+style="[^"]*text-align:\s*(center|right|left|justify)[^"]*">\s*(.*?)\s*</div>', align_replacer, text, flags=re.IGNORECASE | re.DOTALL)

    # --- LIST & TABLE SPACING FIX ---
    # Because of hard_line_breaks, tables and lists merge with text if there's no blank line.
    # This automatically detects missing blank lines and injects them.
    lines = text.split('\n')
    for i in range(1, len(lines)):
        prev_line = lines[i-1].strip()
        curr_line = lines[i].strip()
        
        # If current line starts a table, and previous line is text (not a pipe or empty)
        if curr_line.startswith('|') and prev_line and not prev_line.startswith('|'):
            lines[i] = '\n' + lines[i]
            
        # If current line is a list item, and previous line is text (not empty, not list, not heading)
        is_curr_list = bool(re.match(r'^(\*|-|\+|\d+\.)\s+', curr_line))
        is_prev_list = bool(re.match(r'^(\*|-|\+|\d+\.)\s+', prev_line))
        is_prev_heading = prev_line.startswith('#')
        
        if is_curr_list and prev_line and not is_prev_list and not is_prev_heading:
            lines[i] = '\n' + lines[i]
            
    text = '\n'.join(lines)


    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
        output_path = tmp.name
        
    try:
        # Step 1: Standard One-Step Pandoc Conversion 
        pypandoc.convert_text(
            text, 
            'docx', 
            format='markdown+hard_line_breaks', 
            outputfile=output_path
        )
        
        # Step 2: Native Word Alignment & Page Break Post-Processing
        doc = docx.Document(output_path)
        align_mode = None
        paragraphs_to_delete = []
        
        for p in doc.paragraphs:
            p_text = p.text.strip()
            
            # Detect Page Break Marker and inject native Word page break
            if p_text == "OMNI_PAGE_BREAK":
                p.text = ""
                p.add_run().add_break(WD_BREAK.PAGE)
                continue
            
            # Detect Start Marker
            if p_text.startswith("OMNI_ALIGN_") and p_text.endswith("_START"):
                align_str = p_text.replace("OMNI_ALIGN_", "").replace("_START", "")
                if align_str == "CENTER": align_mode = WD_ALIGN_PARAGRAPH.CENTER
                elif align_str == "RIGHT": align_mode = WD_ALIGN_PARAGRAPH.RIGHT
                elif align_str == "LEFT": align_mode = WD_ALIGN_PARAGRAPH.LEFT
                elif align_str == "JUSTIFY": align_mode = WD_ALIGN_PARAGRAPH.JUSTIFY
                paragraphs_to_delete.append(p)
                continue
                
            # Detect End Marker
            if p_text == "OMNI_ALIGN_END":
                align_mode = None
                paragraphs_to_delete.append(p)
                continue
                
            # Apply True Word Alignment
            if align_mode is not None:
                p.alignment = align_mode

        # Safely delete the invisible markers from the final Word document
        for p in paragraphs_to_delete:
            element = p._element
            element.getparent().remove(element)
            
        doc.save(output_path)
        
        return send_file(output_path, as_attachment=True, download_name='AI_Document.docx')
        
    except Exception as e:
        print(f"Conversion error: {str(e)}")
        return {"error": "Failed to convert text"}, 500
    finally:
        if os.path.exists(output_path):
            os.remove(output_path)

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=10000)
